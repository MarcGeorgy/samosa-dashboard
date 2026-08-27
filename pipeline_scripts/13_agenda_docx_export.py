"""
Builds a Word (.docx) version of the debrief agenda from the same structured
payload 08_debrief_insights.py's build_agenda_payload() produces for the
on-screen/markdown agenda, so the Word doc and the in-app agenda can never
drift apart - this reads the structured dict directly, not the free-form
markdown/Claude-phrased text.

Usage: build_agenda_docx(payload: dict) -> bytes
"""
import io

from docx import Document
from docx.shared import Pt, RGBColor

JPAL_ORANGE = RGBColor(0xE3, 0x59, 0x25)
JPAL_BLUE = RGBColor(0x2D, 0x61, 0x6E)

RA_NOTES_LINES = 10


def _heading(doc, text, level=1, color=JPAL_BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def _bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def build_agenda_docx(payload: dict) -> bytes:
    doc = Document()

    title = doc.add_heading(f"S.A.M.O.S.A Debrief Agenda - {payload['date']}", level=0)
    for run in title.runs:
        run.font.color.rgb = JPAL_ORANGE

    # ------------------------------------------------ questions to ask
    _heading(doc, "Questions to Ask Your Enumerators")
    questions = payload.get("questions_to_ask", [])
    if questions:
        by_enum = {}
        for q in questions:
            by_enum.setdefault(q["enumerator_id"], []).append(q["question"])
        for enumerator_id, qs in by_enum.items():
            p = doc.add_paragraph()
            p.add_run(enumerator_id).bold = True
            for q in qs:
                _bullet(doc, q)
    else:
        doc.add_paragraph("No enumerator meets the check-in threshold this period - "
                          "no targeted questions needed.")

    # ------------------------------------------------ things to monitor
    _heading(doc, "Things to Monitor Yourself")
    for m in payload.get("things_to_monitor", []):
        _bullet(doc, m)

    # ------------------------------------------------ overview
    _heading(doc, "Overview")
    h = payload.get("headline", {})
    doc.add_paragraph(f"{h.get('n_submissions')} submissions, avg quality score {h.get('avg_quality_score')}, "
                      f"{h.get('pct_flagged_any')}% carrying at least one flag.")

    # ------------------------------------------------ check-in / commend
    if payload.get("flag_for_checkin"):
        _heading(doc, "Enumerators to Check In With")
        for r in payload["flag_for_checkin"]:
            _bullet(doc, f"{r['enumerator_id']} (avg {r['avg_quality_score']}): {r['rationale']}")

    if payload.get("commend"):
        _heading(doc, "Enumerators to Commend")
        for r in payload["commend"]:
            _bullet(doc, f"{r['enumerator_id']} (avg {r['avg_quality_score']}): {r['rationale']}")

    # ------------------------------------------------ records to correct
    if payload.get("correction_items"):
        _heading(doc, "Records to Review / Correct on the Call")
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        headers = ["Household", "Enumerator", "Village", "Flags", "Score"]
        for i, col_name in enumerate(headers):
            table.rows[0].cells[i].text = col_name
        for c in payload["correction_items"]:
            row = table.add_row().cells
            row[0].text = str(c["hh_id"])
            row[1].text = str(c["enumerator_id"])
            row[2].text = str(c["village"])
            row[3].text = str(c["flags_triggered"])
            row[4].text = str(c["data_quality_score"])

    # ------------------------------------------------ comment themes
    themes = payload.get("comment_themes", {})
    _heading(doc, "Field Comment Themes")
    if themes.get("overall_summary"):
        doc.add_paragraph(themes["overall_summary"])
    for t in themes.get("themes", []):
        _bullet(doc, f"{t['theme']} ({t['n_mentions']}x, {', '.join(t['affected_enumerators'])}): "
                     f"\"{t['example_quote']}\"")

    # ------------------------------------------------ RA comments (blank)
    _heading(doc, "RA Comments")
    note = doc.add_paragraph()
    note_run = note.add_run("Space for your own notes - fill in during or after the call.")
    note_run.italic = True
    for _ in range(RA_NOTES_LINES):
        p = doc.add_paragraph("_" * 95)
        p.paragraph_format.space_after = Pt(14)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
